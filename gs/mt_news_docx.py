# GS — MT Newswires DOCX Reader
# Architecture: Claude.ai scheduled task fetches MT Newswires via MCP connector,
# writes a DOCX to GroundTruth_v2/. This module reads that DOCX and converts
# articles into standard raw-item dicts for the classification pipeline.
# Also exports a companion JSON file for direct ingestion.
# Last Updated: April 13 2026

import sys
import os
import json
import glob
import re
from datetime import datetime, date

PROJECT_ROOT = r"C:\Users\nagar_7kszmu8\GroundTruth_v2"
if PROJECT_ROOT not in sys.path:
    sys.path.insert(0, PROJECT_ROOT)


# ── KEYWORD-BASED C-TAG MAPPING ──────────────────────────────────────────────
# Matches headline/summary text to assign c_tags and a_tags.
# Mirrors the rawMetadataCodes mapping from the task spec,
# adapted for keyword matching since the DOCX carries text not codes.

CTAG_RULES = [
    # oil / energy / geopolitical
    (r"crude|oil|brent|wti|hormuz|iran|opec|blockade|strait",
     ["C05", "C08"], ["A12", "A15"]),
    # macro / rates / fed
    (r"fomc|fed\b|federal reserve|interest rate|treasury|yield|inflation|gdp|recession|stagflat",
     ["C01", "C02"], ["ALL"]),
    # financing / credit
    (r"bdc|private credit|project finance|infrastructure debt|loan|refinanc|spread|leverage|credit",
     ["C12", "C01"], ["ALL"]),
    # regulatory / ferc / grid
    (r"ferc|interconnection|transmission|grid|regulatory|permit",
     ["C02", "C15"], ["A17"]),
    # solar / renewables / curtailment
    (r"caiso|ercot|curtailment|solar|ppa|negative price|renewable",
     ["C06", "C15"], ["A04", "A05"]),
    # construction / supply chain / tariff
    (r"steel|aluminum|tariff|construction cost|supply chain|commodity",
     ["C09", "C08"], ["ALL"]),
    # data center / digital infra
    (r"data center|hyperscal|gpu|coreweave|ai infrastructure|colocation",
     ["C11", "C14"], ["A09", "A10"]),
    # nuclear / smr
    (r"nuclear|smr|small modular reactor",
     ["C02", "C04"], ["A02", "A25"]),
    # lng
    (r"lng|liquefied natural gas|feedgas",
     ["C05", "C08"], ["A12", "A13"]),
    # equities / general market (low priority — catch-all)
    (r"equit|s&p 500|nasdaq|dow jones|stock market",
     ["C01"], ["ALL"]),
]

# Articles with these keywords are excluded (analyst noise)
EXCLUDE_PATTERNS = re.compile(
    r"analyst.rating.change|price.target.only|insider.trad|social.buzz|"
    r"cryptocurrency|crypto|bitcoin|healthcare|consumer.goods|"
    r"pharma|biotech|weight.loss|drug.approv",
    re.IGNORECASE,
)


def _assign_tags(text: str) -> tuple[list[str], list[str]]:
    """Assign c_tags and a_tags by matching headline+summary text against rules."""
    text_lower = text.lower()
    c_tags = []
    a_tags = []
    for pattern, ctags, atags in CTAG_RULES:
        if re.search(pattern, text_lower):
            for t in ctags:
                if t not in c_tags:
                    c_tags.append(t)
            for t in atags:
                if t not in a_tags:
                    a_tags.append(t)
    return c_tags or ["C01"], a_tags or ["ALL"]


def _find_docx(on_date: str = None) -> str | None:
    """Find the MT Newswires DOCX for the given date (or today)."""
    date_str = on_date or date.today().strftime("%Y-%m-%d")
    path = os.path.join(PROJECT_ROOT, f"MT_Newswires_{date_str}.docx")
    if os.path.isfile(path):
        return path
    # Try glob for any MT_Newswires DOCX with today's date
    pattern = os.path.join(PROJECT_ROOT, f"MT_Newswires*{date_str}*.docx")
    matches = glob.glob(pattern)
    return matches[0] if matches else None


def _parse_docx(path: str, date_str: str) -> list[dict]:
    """Parse the MT Newswires DOCX and return raw article dicts."""
    from docx import Document

    doc = Document(path)
    if not doc.tables:
        return []

    table = doc.tables[0]
    articles = []

    for row in table.rows[1:]:  # skip header row
        cells = [c.text.strip() for c in row.cells]
        if len(cells) < 2:
            continue

        time_et = cells[0]
        full_text = cells[1]
        if not full_text:
            continue

        lines = full_text.split("\n")
        headline = lines[0].strip()
        summary = "\n".join(lines[1:]).strip() if len(lines) > 1 else ""

        if not headline:
            continue

        # Exclude analyst noise
        combined = f"{headline} {summary}"
        if EXCLUDE_PATTERNS.search(combined):
            continue

        c_tags, a_tags = _assign_tags(combined)

        pub_date = f"{date_str} {time_et}" if time_et else date_str

        articles.append({
            "headline":         headline[:120],
            "summary":          summary[:500],
            "url":              "",
            "source_name":      "MT Newswires",
            "publication_date": pub_date,
            "c_tags":           c_tags,
            "a_tags":           a_tags,
            "raw_content":      combined[:3000],
        })

    return articles


def _write_json(articles: list[dict], date_str: str) -> str:
    """Write articles to a companion JSON file. Returns the path."""
    path = os.path.join(PROJECT_ROOT, f"MT_Newswires_{date_str}.json")
    with open(path, "w", encoding="utf-8") as f:
        json.dump(articles, f, indent=2, ensure_ascii=False)
    return path


def generate(on_date: str = None) -> list[dict]:
    """
    Main entry point. Read today's MT Newswires DOCX, parse articles,
    write companion JSON, return raw article dicts for the pipeline.

    Args:
        on_date: Date string YYYY-MM-DD. Defaults to today.

    Returns:
        List of raw article dicts matching fetch_rss() output format.
        Returns empty list if no DOCX found for the date.
    """
    date_str = on_date or date.today().strftime("%Y-%m-%d")

    path = _find_docx(date_str)
    if not path:
        return []

    articles = _parse_docx(path, date_str)
    if not articles:
        return []

    json_path = _write_json(articles, date_str)
    print(f"  MT Newswires DOCX: {len(articles)} articles from {os.path.basename(path)}")
    print(f"  JSON export: {os.path.basename(json_path)}")

    return articles


# ── TEST ─────────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    print("=" * 55)
    print("MT NEWSWIRES DOCX — READER TEST")
    print("=" * 55)

    items = generate()
    print(f"\nArticles parsed: {len(items)}")

    if items:
        print("\nArticles:")
        for i, item in enumerate(items, 1):
            print(f"  {i}. [{', '.join(item['c_tags'])}] {item['headline'][:70]}")
            print(f"     pub={item['publication_date']}")
    else:
        today = date.today().strftime("%Y-%m-%d")
        print(f"\n  No DOCX found: MT_Newswires_{today}.docx")
        print("  Place the file in GroundTruth_v2/ and re-run.")

    print("\ngs/mt_news_docx.py operational.")
